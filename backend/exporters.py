"""
Export the tailored resume (Markdown) to PDF and DOCX, replicating the exact
formatting of the reference resume (Purpose Financial):

  Page       : US Letter, margins  L/R 0.40"  T/B 0.50"
  Font       : Arial (embedded from Windows fonts; Helvetica fallback)
  Color      : 100% black
  Leading    : 13 pt
  Name+Title : 12 pt bold, one line  "Name • Title"
  Contact    : 11 pt regular
  Section hdr: 11 pt bold, UPPERCASE + ':', with a 0.8 pt full-width rule under it
  Body/bullet: 11 pt regular, bullet = •
"""
import io
import os
import re
from html import escape

BLACK = "#000000"

# ---- sizes / spacing (points) ---------------------------------------------
BODY = 11.0
NAME = 13.0  # name + title = body + 2, bold black
LEADING = 13.0
MARGIN_LR = 0.40 * 72
MARGIN_TB = 0.50 * 72
RULE_WIDTH = 0.8

# ---------------------------------------------------------------------------
# Font registration (Arial embedded; falls back to Helvetica core)
# ---------------------------------------------------------------------------
_FONT = "Helvetica"
_FONT_B = "Helvetica-Bold"
_fonts_ready = False


def _ensure_fonts():
    global _FONT, _FONT_B, _fonts_ready
    if _fonts_ready:
        return
    _fonts_ready = True
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        win = os.environ.get("WINDIR", "C:/Windows")
        bundled = os.path.join(os.path.dirname(__file__), "fonts")
        # Try, in order: bundled TTFs (works on Linux/Railway), then Windows Arial.
        candidates = [
            (os.path.join(bundled, "Arial.ttf"), os.path.join(bundled, "Arial-Bold.ttf")),
            (os.path.join(bundled, "LiberationSans-Regular.ttf"),
             os.path.join(bundled, "LiberationSans-Bold.ttf")),
            (os.path.join(win, "Fonts", "arial.ttf"), os.path.join(win, "Fonts", "arialbd.ttf")),
        ]
        for reg, bold in candidates:
            if os.path.exists(reg) and os.path.exists(bold):
                pdfmetrics.registerFont(TTFont("Arial", reg))
                pdfmetrics.registerFont(TTFont("Arial-Bold", bold))
                # Map the family so inline <b> markup resolves to bold.
                pdfmetrics.registerFontFamily(
                    "Arial", normal="Arial", bold="Arial-Bold",
                    italic="Arial", boldItalic="Arial-Bold",
                )
                _FONT, _FONT_B = "Arial", "Arial-Bold"
                break
    except Exception:  # noqa: BLE001 — keep Helvetica fallback
        pass


# ---------------------------------------------------------------------------
# Markdown -> typed elements
# ---------------------------------------------------------------------------
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_BOLD_LINE = re.compile(r"^\*\*(.+?)\*\*[\s,–—-]*$")
# per-job "Technologies Used:" line (with or without a leading bullet / ** wrap)
_TECHLABEL = re.compile(r"^[-*]?\s*\*{0,2}(technolog|tech stack|tools used|environment)", re.I)
_DATE = re.compile(r"((?:19|20)\d{2}|present)", re.I)
_CONTACT = re.compile(r"[@|]|linkedin|github|https?://|\(\d{3}\)|\d{3}[.\-]\d{3}[.\-]\d{4}", re.I)


def _clean_contact(text: str) -> str:
    """Phone first, email second, drop city/state and everything else."""
    parts = [p.strip() for p in re.split(r"[|•]", text) if p.strip()]
    phone = next((p for p in parts if re.search(r"\d{3}[)\s.\-]*\d{3}[\s.\-]*\d{4}", p)), None)
    email = next((p for p in parts if "@" in p), None)
    keep = [x for x in (phone, email) if x]
    return " | ".join(keep) if keep else text


def _split_jobline(text: str):
    """Split a job-title line into (left, right-date). Date = the | field with a
    year/Present. Returns (left, '') if no date found."""
    fields = [f.strip() for f in text.split("|")
              if not re.fullmatch(r"(location\s*)?(not listed|not specified|n/?a|none|unknown|tbd)",
                                  f.strip(), re.I)]
    for i, f in enumerate(fields):
        if _DATE.search(f):
            date = fields[i]
            left = " | ".join(fields[:i] + fields[i + 1:])
            return left.strip(" |"), date
    return text, ""


def _classify(md: str):
    out = []
    name_done = False
    section_seen = False
    in_skills = False
    prev_kind = None
    for raw in md.replace("\r\n", "\n").split("\n"):
        s = raw.strip()
        if not s:
            prev_kind = None
            continue
        if re.fullmatch(r"[-*_]{3,}", s):
            continue  # skip markdown horizontal rules (--- between jobs)
        if s.startswith(">") or re.search(r"see above|consolidated under", s, re.I):
            continue  # skip LLM consolidation stubs
        if s.startswith("## ") or s.startswith("### "):
            label = s.lstrip("#").strip()
            in_skills = "skill" in label.lower()
            out.append(("section", label))
            section_seen = True
        elif s.startswith("# "):
            out.append(("name", s[2:].strip()))
            name_done = True
        elif _TECHLABEL.match(s):
            out.append(("tech", re.sub(r"^[-*]\s+", "", s).replace("*", "").strip()))
        elif s.startswith(("- ", "* ")):
            txt = s[2:].replace("*", "").strip()
            out.append(("skillbullet" if in_skills else "bullet", txt))
        elif _BOLD_LINE.match(s):
            kind = "jobtitle" if section_seen else "headline"
            out.append((kind, _BOLD_LINE.match(s).group(1).strip()))
        elif not section_seen and not name_done:
            out.append(("name", s))
            name_done = True
        elif prev_kind == "name" and _CONTACT.search(s):
            out.append(("contact", s))
        else:
            out.append(("body", s))
        prev_kind = out[-1][0] if out else None
    return out


def _label_markup(text: str) -> str:
    """Bold the label up to the first colon, regular the rest (reportlab markup)."""
    text = text.replace("*", "").strip()
    if ":" in text:
        label, rest = text.split(":", 1)
        return f"<b>{escape(label)}:</b>{escape(rest)}"
    return escape(text)


def _inline(text: str) -> str:
    parts, pos = [], 0
    for m in _BOLD.finditer(text):
        parts.append(escape(text[pos:m.start()]))
        parts.append("<b>" + escape(m.group(1)) + "</b>")
        pos = m.end()
    parts.append(escape(text[pos:]))
    return "".join(parts)


_CONTENT_KINDS = {"bullet", "skillbullet", "tech", "body"}


def _prune_empty(elements):
    """Drop bulletless job stubs and then empty section headers (kills the
    duplicate 'Molina (Earlier Engagement)' with no bullets and empty CERTIFICATIONS)."""
    # Pass 1: in Experience/Projects, a job title needs at least one bullet before
    # the next title/section. (Education/Certifications titles are fine without one.)
    keep = [True] * len(elements)
    cur_sec = ""
    for i, (kind, text) in enumerate(elements):
        if kind == "section":
            cur_sec = text.lower()
            continue
        if kind != "jobtitle" or not ("experience" in cur_sec or "project" in cur_sec):
            continue
        has = False
        for j in range(i + 1, len(elements)):
            k2 = elements[j][0]
            if k2 in ("jobtitle", "section"):
                break
            if k2 in _CONTENT_KINDS:
                has = True
                break
        if not has:
            keep[i] = False
    elements = [e for e, k in zip(elements, keep) if k]

    # Pass 2: a section needs ANY element (incl. a job/degree title) before the next section.
    keep = [True] * len(elements)
    for i, (kind, _) in enumerate(elements):
        if kind != "section":
            continue
        nxt = elements[i + 1][0] if i + 1 < len(elements) else "section"
        if nxt == "section":
            keep[i] = False
    return [e for e, k in zip(elements, keep) if k]


def _split_header(elements):
    name = contact = headline = None
    rest = []
    for kind, text in elements:
        if kind == "name" and name is None:
            name = text
        elif kind == "contact" and contact is None:
            contact = text
        elif kind == "headline" and headline is None:
            headline = text
        else:
            rest.append((kind, text))
    return name, contact, headline, rest


# ---------------------------------------------------------------------------
# PDF — auto-fit: uniformly shrink type (never truncate) to cap at MAX_PAGES
# ---------------------------------------------------------------------------
MAX_PAGES = 2
# Uniform shrink steps (font + leading + spacing). Margins nudge gently.
_SCALES = [1.0, 0.96, 0.92, 0.88, 0.84, 0.80]


def _render_pdf(elements, scale):
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Table, TableStyle

    black = HexColor(BLACK)
    body = BODY * scale
    name_sz = body + 2
    lead = LEADING * scale
    mlr = MARGIN_LR * (0.9 + 0.1 * scale)
    mtb = MARGIN_TB * (0.9 + 0.1 * scale)
    usable_w = letter[0] - 2 * mlr
    sp = scale  # spacing multiplier

    def ps(nm, size, bold=False, **kw):
        params = {"fontName": _FONT_B if bold else _FONT, "fontSize": size,
                  "leading": lead, "textColor": black, "alignment": TA_LEFT}
        params.update(kw)
        return ParagraphStyle(nm, **params)

    st_name = ps("name", name_sz, bold=True, leading=name_sz + 2, spaceAfter=1)
    st_contact = ps("contact", body, spaceAfter=4 * sp)
    st_section = ps("section", body + 1, bold=True, spaceBefore=8 * sp, spaceAfter=2 * sp)
    st_job = ps("job", body, bold=True, spaceBefore=4 * sp, spaceAfter=1)
    st_job_r = ps("jobr", body, bold=True, alignment=TA_RIGHT, spaceBefore=4 * sp, spaceAfter=1)
    st_body = ps("body", body, spaceAfter=2 * sp)
    st_tech = ps("tech", body, spaceAfter=3 * sp)
    st_bullet = ps("bullet", body, leftIndent=13, bulletIndent=2, spaceAfter=2 * sp)

    name, contact, headline, rest = _split_header(elements)
    story = []
    if name or headline:
        head = "<b>" + _inline(name or "") + "</b>"
        if headline:
            head += "&nbsp;&nbsp;<b>-</b>&nbsp;&nbsp;<b>" + _inline(headline) + "</b>"
        story.append(Paragraph(head, st_name))
    if contact:
        story.append(Paragraph(escape(_clean_contact(contact)), st_contact))

    for kind, text in rest:
        if kind == "section":
            story.append(Paragraph(escape(text).upper().rstrip(":") + ":", st_section))
            story.append(HRFlowable(width="100%", thickness=RULE_WIDTH,
                                    color=black, spaceBefore=1, spaceAfter=4 * sp))
        elif kind == "jobtitle":
            left, date = _split_jobline(text)
            if date:
                lp = Paragraph("<b>" + _inline(left) + "</b>", st_job)
                rp = Paragraph("<b>" + escape(date) + "</b>", st_job_r)
                t = Table([[lp, rp]], colWidths=[usable_w * 0.72, usable_w * 0.28])
                t.hAlign = "LEFT"  # align job title with section headers, not frame edge
                t.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 3 * sp),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
                story.append(t)
            else:
                story.append(Paragraph("<b>" + _inline(text) + "</b>", st_job))
        elif kind == "tech":
            story.append(Paragraph(_label_markup(text), st_tech))
        elif kind == "skillbullet":
            story.append(Paragraph(_label_markup(text), st_bullet, bulletText="•"))
        elif kind == "bullet":
            story.append(Paragraph(_inline(text), st_bullet, bulletText="•"))
        else:
            story.append(Paragraph(_inline(text), st_body))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=mlr, rightMargin=mlr,
                            topMargin=mtb, bottomMargin=mtb, title="Tailored Resume")
    doc.build(story)
    return buf.getvalue(), doc.page


def _choose_scale(elements):
    """Smallest shrink that fits MAX_PAGES; else the tightest. Returns (scale, bytes)."""
    last = None
    for scale in _SCALES:
        data, pages = _render_pdf(elements, scale)
        if pages <= MAX_PAGES:
            return scale, data
        last = (scale, data)
    return last  # tightest attempt


def to_pdf(md: str) -> bytes:
    _ensure_fonts()
    return _choose_scale(_prune_empty(_classify(md)))[1]


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _para_bottom_border(paragraph):
    """Add a thin bottom border (the section rule) to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")      # 0.75 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def to_docx(md: str) -> bytes:
    _ensure_fonts()
    import docx
    from docx.shared import Inches, Pt, RGBColor

    elements = _prune_empty(_classify(md))
    scale = _choose_scale(elements)[0]  # same auto-fit scale as the PDF
    body = BODY * scale
    name_sz = body + 2
    lead = LEADING * scale

    black = RGBColor(0, 0, 0)
    doc = docx.Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.40)
    sec.top_margin = sec.bottom_margin = Inches(0.50)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(body)
    normal.paragraph_format.space_after = Pt(2 * scale)
    normal.paragraph_format.line_spacing = Pt(lead)

    def runs(p, text, size, bold=False):
        for chunk, b in _split_bold(text):
            r = p.add_run(chunk)
            r.font.name = "Arial"
            r.font.size = Pt(size)
            r.bold = bold or b
            r.font.color.rgb = black

    name, contact, headline, rest = _split_header(elements)

    if name or headline:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        runs(p, name or "", name_sz, bold=True)
        if headline:
            runs(p, "     -     " + headline, name_sz, bold=True)
    if contact:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4 * scale)
        runs(p, _clean_contact(contact), body)

    for kind, text in rest:
        if kind == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8 * scale)
            p.paragraph_format.space_after = Pt(3 * scale)
            runs(p, text.upper().rstrip(":") + ":", body + 1, bold=True)
            _para_bottom_border(p)
        elif kind == "jobtitle":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 * scale)
            left, date = _split_jobline(text)
            if date:
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Inches as _In
                right_pos = _In(8.5 - 0.40 - 0.40)  # page width minus L/R margins
                p.paragraph_format.tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)
                runs(p, left, body, bold=True)
                runs(p, "\t" + date, body, bold=True)
            else:
                runs(p, text, body, bold=True)
        elif kind in ("tech", "skillbullet"):
            p = (doc.add_paragraph(style="List Bullet") if kind == "skillbullet"
                 else doc.add_paragraph())
            p.paragraph_format.space_after = Pt((2 if kind == "skillbullet" else 3) * scale)
            t = text.replace("*", "")
            if ":" in t:
                label, restt = t.split(":", 1)
                runs(p, label + ":", body, bold=True)
                runs(p, restt, body)
            else:
                runs(p, t, body)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2 * scale)
            runs(p, text, body)
        else:
            runs(doc.add_paragraph(), text, body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _split_bold(text: str):
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], False
        yield m.group(1), True
        pos = m.end()
    if pos < len(text):
        yield text[pos:], False
